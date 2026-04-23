from pathlib import Path

from loaders import doc_loader


def test_extract_text_from_txt(tmp_path: Path):
    f = tmp_path / "t.txt"
    f.write_text("hello world", encoding="utf-8")
    text, error = doc_loader.extract_text_from_file(f)
    assert text == "hello world"
    assert error is None


def test_extract_text_from_unsupported(tmp_path: Path):
    f = tmp_path / "weird.xyz"
    f.write_text("content", encoding="utf-8")
    text, error = doc_loader.extract_text_from_file(f)
    assert text == ""
    assert "Unsupported" in error


def test_extract_text_from_docx(tmp_path: Path):
    from docx import Document
    f = tmp_path / "sample.docx"
    doc = Document()
    doc.add_paragraph("First paragraph.")
    doc.add_paragraph("Second paragraph.")
    doc.save(str(f))
    text, error = doc_loader.extract_text_from_file(f)
    assert error is None
    assert "First paragraph." in text
    assert "Second paragraph." in text


def test_load_institutional_docs_missing_folder(tmp_path: Path):
    assert doc_loader.load_institutional_docs(str(tmp_path / "nope")) == []


def test_load_institutional_docs_skips_unsupported(tmp_path: Path):
    folder = tmp_path / "d"
    folder.mkdir()
    (folder / "a.txt").write_text("alpha", encoding="utf-8")
    (folder / "skip.xyz").write_text("ignored", encoding="utf-8")
    docs = doc_loader.load_institutional_docs(str(folder))
    names = [d["filename"] for d in docs]
    assert "a.txt" in names
    assert "skip.xyz" not in names
    a_entry = next(d for d in docs if d["filename"] == "a.txt")
    assert a_entry["text"] == "alpha"
    assert a_entry["char_count"] == 5
    assert a_entry["error"] is None


def test_load_institutional_docs_recurses(tmp_path: Path):
    nested = tmp_path / "d" / "sub"
    nested.mkdir(parents=True)
    (nested / "deep.txt").write_text("deep", encoding="utf-8")
    docs = doc_loader.load_institutional_docs(str(tmp_path / "d"))
    assert any(d["filename"] == "deep.txt" for d in docs)
